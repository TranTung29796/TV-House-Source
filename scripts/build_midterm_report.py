from __future__ import annotations

import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "report-assets"
OUTPUT = DOCS / "BAO_CAO_GIUA_KY_TV_HOUSE.docx"

NAVY = "14213D"
RED = "D7264D"
TEAL = "087F8C"
BLUE = "2E74B5"
LIGHT = "F2F4F7"
MID = "667085"
WHITE = "FFFFFF"
BLACK = "171A1F"

ASSETS.mkdir(parents=True, exist_ok=True)


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_run(run, *, size=10.5, bold=False, color=BLACK, italic=False, font="Arial"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)
    return run


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, fld_sep, text, fld_end])


def make_code_image(source_path: Path, start_marker: str, end_marker: str, output: Path, title: str) -> None:
    lines = source_path.read_text(encoding="utf-8").splitlines()
    start = next(i for i, line in enumerate(lines) if start_marker in line)
    end = next(i for i, line in enumerate(lines[start:], start) if end_marker in line)
    snippet = lines[start : end + 1]
    font = ImageFont.truetype("/System/Library/Fonts/SFNSMono.ttf", 26)
    title_font = ImageFont.truetype("/System/Library/Fonts/SFNSMono.ttf", 25)
    width = 1700
    line_height = 39
    height = 95 + line_height * len(snippet) + 35
    image = Image.new("RGB", (width, height), "#111827")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((18, 18, width - 18, height - 18), radius=18, fill="#111827", outline="#334155", width=2)
    draw.ellipse((42, 37, 60, 55), fill="#fb7185")
    draw.ellipse((69, 37, 87, 55), fill="#fbbf24")
    draw.ellipse((96, 37, 114, 55), fill="#34d399")
    draw.text((135, 31), title, font=title_font, fill="#e5e7eb")
    keyword = re.compile(r"\b(export|function|const|if|for|try|catch|return|throw|new|await|async|CREATE|TABLE|FOREIGN|KEY|REFERENCES|INSERT|INTO|VALUES|UPDATE|WHERE)\b")
    y = 82
    for index, line in enumerate(snippet, start=start + 1):
        draw.text((35, y), f"{index:>3}", font=font, fill="#64748b")
        x = 125
        parts = keyword.split(line.expandtabs(2))
        for part in parts:
            color = "#f472b6" if keyword.fullmatch(part) else "#dbeafe"
            draw.text((x, y), part, font=font, fill=color)
            x += draw.textlength(part, font=font)
        y += line_height
    image.save(output)


def make_architecture_image(output: Path) -> None:
    image = Image.new("RGB", (1600, 760), "white")
    draw = ImageDraw.Draw(image)
    title_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 42)
    font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 29)
    small = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 23)
    draw.text((70, 45), "Kiến trúc tổng thể TV House", font=title_font, fill="#14213d")
    boxes = [
        (90, 200, 420, 470, "TRÌNH DUYỆT", "Trang chủ\nGiỏ hàng\nQuản trị", "#E8F4F5"),
        (635, 200, 965, 470, "NEXT.JS", "React UI\nAPI Routes\nKiểm tra dữ liệu", "#FDECEF"),
        (1180, 200, 1510, 470, "SQLITE", "products\norders\norder_items", "#EEF2FF"),
    ]
    for x1, y1, x2, y2, heading, body, fill in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=16, fill=fill, outline="#cbd5e1", width=3)
        draw.text((x1 + 34, y1 + 35), heading, font=font, fill="#d7264d")
        draw.multiline_text((x1 + 34, y1 + 105), body, font=font, fill="#14213d", spacing=16)
    for x1, x2 in ((420, 635), (965, 1180)):
        draw.line((x1 + 30, 335, x2 - 30, 335), fill="#087f8c", width=8)
        draw.polygon([(x2 - 48, 316), (x2 - 20, 335), (x2 - 48, 354)], fill="#087f8c")
    draw.text((112, 560), "HTTP/JSON", font=small, fill="#667085")
    draw.text((650, 560), "Prepared statement + transaction", font=small, fill="#667085")
    draw.text((90, 650), "Luồng dữ liệu: thao tác người dùng → xử lý nghiệp vụ → lưu trữ bền vững", font=font, fill="#14213d")
    image.save(output)


def crop_report_images() -> None:
    crops = [
        (DOCS / "screenshots/01-trang-chu.png", ASSETS / "result-home.png", 0, 0, 1425, 980),
        (DOCS / "screenshots/02-gio-hang.png", ASSETS / "result-cart.png", 0, 0, 1425, 1030),
    ]
    for src, dst, x1, y1, x2, y2 in crops:
        image = Image.open(src).convert("RGB")
        image.crop((x1, y1, min(x2, image.width), min(y2, image.height))).save(dst, quality=95)


make_code_image(ROOT / "src/lib/store-db.ts", "CREATE TABLE IF NOT EXISTS products", ") STRICT;", ASSETS / "code-schema.png", "src/lib/store-db.ts — Khởi tạo bảng products")
make_code_image(ROOT / "src/lib/store-db.ts", "db.exec(\"BEGIN IMMEDIATE\")", "db.exec(\"ROLLBACK\")", ASSETS / "code-transaction.png", "src/lib/store-db.ts — Transaction tạo hóa đơn")
make_code_image(ROOT / "src/app/api/orders/route.ts", "export async function POST", "return NextResponse.json({ order }", ASSETS / "code-api-order.png", "src/app/api/orders/route.ts — API đặt hàng")
make_architecture_image(ASSETS / "architecture.png")
crop_report_images()

doc = Document()
doc.settings.odd_and_even_pages_header_footer = False
section = doc.sections[0]
section.different_first_page_header_footer = False
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(10.5)
normal.font.color.rgb = rgb(BLACK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

for name, size, before, after, color in (
    ("Heading 1", 16, 16, 8, BLUE),
    ("Heading 2", 13, 12, 6, BLUE),
    ("Heading 3", 12, 8, 4, NAVY),
):
    style = styles[name]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = rgb(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_run(header.add_run("TV HOUSE  |  BÁO CÁO GIỮA KỲ"), size=8.5, bold=True, color=MID)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(footer.add_run("Trang "), size=8.5, color=MID)
add_page_field(footer)


current_num_id = None


def create_numbering_sequence() -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    level.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "space")
    level.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def start_page(title: str, subtitle: str | None = None, *, first=False):
    global current_num_id
    current_num_id = create_numbering_sequence()
    p = doc.add_paragraph()
    if not first:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.space_after = Pt(12)
    set_run(p.add_run(title.upper()), size=18, bold=True, color=NAVY)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), RED)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(14)
        set_run(p.add_run(subtitle), size=11, italic=True, color=MID)


def para(text: str, *, bold_prefix: str | None = None, align=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        set_run(p.add_run(bold_prefix), bold=True)
        set_run(p.add_run(text[len(bold_prefix):]))
    else:
        set_run(p.add_run(text))
    return p


def bullet(text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    set_run(p.add_run(text))
    return p


def numbered(text: str):
    global current_num_id
    if current_num_id is None:
        current_num_id = create_numbering_sequence()
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(current_num_id))
    num_pr.extend([ilvl, num_id])
    p_pr.append(num_pr)
    set_run(p.add_run(text))
    return p


def set_table_geometry(table, widths):
    width_twips = [round(width * 1440) for width in widths]
    table_pr = table._tbl.tblPr
    layout = table_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    table_width = table_pr.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:w"), str(sum(width_twips)))
    table_width.set(qn("w:type"), "dxa")
    indent = table_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        table_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in width_twips:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, width_twips):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.append(tc_width)
            tc_width.set(qn("w:w"), str(width))
            tc_width.set(qn("w:type"), "dxa")


def add_table(headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.style = "Table Grid"
    for i, (header_text, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[i]
        cell.width = Inches(width)
        set_cell_shading(cell, LIGHT)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(cell.paragraphs[0].add_run(header_text), size=9, bold=True, color=NAVY)
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for i, (value, width) in enumerate(zip(row_data, widths)):
            cells[i].width = Inches(width)
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(str(value)), size=8.7)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_picture(path: Path, width=6.3, caption: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    picture = p.add_run().add_picture(str(path), width=Inches(width))
    alt_text = caption or f"Hình minh họa {path.stem.replace('-', ' ')}"
    picture._inline.docPr.set("descr", alt_text)
    picture._inline.docPr.set("title", alt_text)
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(7)
        set_run(cp.add_run(caption), size=8.5, italic=True, color=MID)


# 1. Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(62)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(p.add_run("TRƯỜNG ĐẠI HỌC / CAO ĐẲNG"), size=12, bold=True, color=NAVY)
para("KHOA CÔNG NGHỆ THÔNG TIN", align=WD_ALIGN_PARAGRAPH.CENTER)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(70)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(p.add_run("BÁO CÁO BÀI KIỂM TRA GIỮA KỲ"), size=15, bold=True, color=RED)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
p.paragraph_format.space_after = Pt(12)
set_run(p.add_run("XÂY DỰNG WEBSITE BÁN TIVI"), size=27, bold=True, color=NAVY)
para("TV HOUSE", align=WD_ALIGN_PARAGRAPH.CENTER)
add_picture(ROOT / "public/images/tv-hero.png", width=5.7)
meta = doc.add_table(rows=4, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta.autofit = False
for row, values in enumerate((("Sinh viên", "[Điền họ và tên]"), ("Mã sinh viên", "[Điền mã sinh viên]"), ("Lớp", "[Điền lớp học phần]"), ("Giảng viên", "[Điền tên giảng viên]"))):
    for col, value in enumerate(values):
        cell = meta.cell(row, col)
        cell.width = Inches(1.55 if col == 0 else 3.65)
        set_cell_margins(cell, 70, 100, 70, 100)
        set_run(cell.paragraphs[0].add_run(value), size=10, bold=col == 0, color=NAVY if col == 0 else BLACK)
set_repeat_table_header(meta.rows[0])
para("Đà Nẵng, tháng 8 năm 2026", align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

# 2
start_page("Thông tin bài làm", "Phiếu tóm tắt để giảng viên đối chiếu nhanh")
add_table(["Nội dung", "Thông tin"], [
    ("Tên đề tài", "Xây dựng website bán tivi TV House"),
    ("Thời gian", "01 tuần kể từ ngày giao đề"),
    ("Công nghệ", "Next.js 15, React 19, TypeScript, SQLite"),
    ("Cơ sở dữ liệu", "3 bảng: products, orders, order_items"),
    ("Chức năng", "Trang chủ, giỏ hàng, đặt hàng, CRUD sản phẩm, quản lý hóa đơn"),
    ("Mã nguồn", "Thư mục web-runtime và gói TV-House-Source.zip"),
    ("Cách chạy", "Xem README.md và docs/HUONG_DAN_TRIEN_KHAI.md"),
], [1.65, 4.85])
doc.add_heading("Cam kết", level=2)
para("Báo cáo mô tả đúng chương trình được xây dựng trong thư mục mã nguồn. Ảnh giao diện, lược đồ dữ liệu và các đoạn mã minh họa đều được lấy từ phiên bản đã chạy kiểm thử.")
doc.add_heading("Thông tin cần hoàn thiện trước khi nộp", level=2)
bullet("Điền họ tên, mã sinh viên, lớp học phần và tên giảng viên ở trang bìa.")
bullet("Tải gói mã nguồn lên GitHub hoặc Google Drive và điền liên kết chia sẻ vào trang 28.")
bullet("Đổi tên file báo cáo theo quy định của lớp nếu giảng viên có yêu cầu riêng.")

# 3
start_page("Tóm tắt đề tài", "Tổng quan kết quả thực hiện")
para("TV House là website bán tivi có giao diện dành cho khách hàng và khu vực quản trị dành cho nhân viên cửa hàng. Khách hàng có thể tìm kiếm sản phẩm, lọc theo thương hiệu, thêm sản phẩm vào giỏ, điều chỉnh số lượng và gửi thông tin đặt hàng. Khi xác nhận, hệ thống tạo hóa đơn cùng các dòng chi tiết trong cơ sở dữ liệu SQLite.")
para("Khu vực quản trị hỗ trợ thêm, sửa, xóa sản phẩm; theo dõi giá bán, tồn kho; xem danh sách hóa đơn và cập nhật trạng thái. Mỗi chi tiết hóa đơn lưu tên và đơn giá tại thời điểm mua để dữ liệu lịch sử không bị thay đổi khi sản phẩm được cập nhật sau này.")
add_table(["Tiêu chí", "Kết quả"], [
    ("Đúng yêu cầu 3 bảng", "Đạt"), ("Trang home", "Đạt"), ("Trang giỏ hàng", "Đạt"),
    ("CRUD sản phẩm", "Đạt"), ("Hóa đơn và chi tiết", "Đạt"), ("Chuyển máy khác", "Có lockfile, Node/pnpm cố định và hướng dẫn"),
], [3.0, 3.5])
doc.add_heading("Kết quả nổi bật", level=2)
bullet("8 sản phẩm mẫu được tạo tự động khi chạy lần đầu.")
bullet("Tạo hóa đơn theo transaction và tự động trừ tồn kho.")
bullet("Responsive, không tràn ngang ở kích thước điện thoại 390 px.")
bullet("Build production, lint và typecheck đều hoàn thành thành công.")

# 4
start_page("Mục lục", "Các nội dung chính trong báo cáo")
toc_rows = [
    ("1", "Giới thiệu và mục tiêu", "5–6"), ("2", "Phân tích yêu cầu và kế hoạch", "7–8"),
    ("3", "Công nghệ và kiến trúc", "9–11"), ("4", "Thiết kế cơ sở dữ liệu", "12–14"),
    ("5", "Thiết kế giao diện và chức năng", "15–18"), ("6", "Thiết kế API", "19"),
    ("7", "Các đoạn mã chính", "20–22"), ("8", "Kiểm thử và kết quả chạy", "23–26"),
    ("9", "Triển khai, kết luận và tài liệu", "27–29"),
]
add_table(["Chương", "Nội dung", "Trang"], toc_rows, [0.75, 4.85, 0.9])
para("Lưu ý: số trang được đánh tự động ở chân trang. Nếu chỉnh sửa nội dung trong Microsoft Word, cần kiểm tra lại mục lục trước khi in hoặc xuất PDF.", after=0)

# 5
start_page("1. Giới thiệu", "Bối cảnh và lý do chọn đề tài")
para("Thị trường điện tử tiêu dùng có số lượng mẫu TV lớn, khác nhau về thương hiệu, kích thước, công nghệ màn hình và giá bán. Một website bán hàng cần trình bày thông tin dễ so sánh, hỗ trợ đặt hàng nhanh và giúp nhân viên cập nhật danh mục mà không chỉnh sửa trực tiếp mã nguồn.")
para("Đề tài TV House mô phỏng quy trình cơ bản của một cửa hàng: trưng bày sản phẩm, lập giỏ hàng, tiếp nhận thông tin giao hàng, tạo hóa đơn và quản trị dữ liệu. Đây là phạm vi phù hợp cho bài giữa kỳ vì kết hợp được thiết kế giao diện, lập trình phía máy chủ và cơ sở dữ liệu quan hệ.")
doc.add_heading("1.1 Đối tượng sử dụng", level=2)
bullet("Khách hàng: xem sản phẩm, quản lý giỏ hàng và đặt mua TV.")
bullet("Nhân viên cửa hàng: quản lý danh mục, tồn kho và tình trạng hóa đơn.")
doc.add_heading("1.2 Giá trị của hệ thống", level=2)
bullet("Thông tin sản phẩm tập trung và có thể tìm kiếm nhanh.")
bullet("Hạn chế sai lệch tổng tiền nhờ tính toán lại ở phía máy chủ.")
bullet("Dữ liệu nằm trong một file SQLite dễ sao lưu và chuyển máy.")

# 6
start_page("2. Mục tiêu và phạm vi", "Tiêu chí hoàn thành bài giữa kỳ")
doc.add_heading("2.1 Mục tiêu chức năng", level=2)
numbered("Xây dựng trang home có hình ảnh, sản phẩm, tìm kiếm và bộ lọc.")
numbered("Xây dựng trang giỏ hàng có tăng/giảm số lượng, xóa và thanh toán.")
numbered("Xây dựng chức năng thêm, sửa, xóa sản phẩm thông qua trang quản trị.")
numbered("Lưu hóa đơn và chi tiết hóa đơn vào SQLite, đồng thời cập nhật tồn kho.")
numbered("Cung cấp tài liệu để project chạy được trên máy tính khác.")
doc.add_heading("2.2 Phạm vi", level=2)
para("Hệ thống tập trung vào luồng bán hàng cơ bản. Thanh toán trực tuyến, đăng nhập phân quyền, tải ảnh lên máy chủ và quản lý vận chuyển thực tế chưa nằm trong phạm vi một tuần. Phương thức chuyển khoản hiện được ghi nhận ở hóa đơn, chưa kết nối cổng ngân hàng.")
doc.add_heading("2.3 Tiêu chí chấp nhận", level=2)
add_table(["Mã", "Tiêu chí", "Mức ưu tiên"], [("AC-01", "Có đúng ba bảng dữ liệu và khóa ngoại", "Cao"), ("AC-02", "CRUD sản phẩm hoạt động", "Cao"), ("AC-03", "Đặt hàng tạo đủ hóa đơn và chi tiết", "Cao"), ("AC-04", "Giao diện dùng được trên mobile", "Trung bình")], [0.8, 4.6, 1.1])

# 7
start_page("3. Phân tích yêu cầu", "Yêu cầu chức năng và phi chức năng")
doc.add_heading("3.1 Yêu cầu chức năng", level=2)
add_table(["Mã", "Mô tả", "Tác nhân"], [
    ("FR-01", "Xem, tìm kiếm và lọc sản phẩm", "Khách"), ("FR-02", "Thêm và chỉnh số lượng trong giỏ", "Khách"),
    ("FR-03", "Nhập thông tin và tạo hóa đơn", "Khách"), ("FR-04", "Thêm/sửa/xóa sản phẩm", "Nhân viên"),
    ("FR-05", "Xem chi tiết, đổi trạng thái hóa đơn", "Nhân viên"),
], [0.8, 4.5, 1.2])
doc.add_heading("3.2 Yêu cầu phi chức năng", level=2)
bullet("Tính toàn vẹn: tạo hóa đơn và trừ tồn kho trong một transaction.")
bullet("Khả năng mang theo: không phụ thuộc MySQL/SQL Server cài ngoài.")
bullet("Khả năng sử dụng: nút, biểu mẫu và thông báo rõ ràng bằng tiếng Việt.")
bullet("Hiệu năng: thêm index cho truy vấn thương hiệu, ngày hóa đơn và khóa ngoại.")
bullet("Tương thích: cố định Node.js 22, pnpm 9.12 và lockfile thư viện.")

# 8
start_page("4. Kế hoạch thực hiện", "Lịch phát triển trong một tuần")
add_table(["Ngày", "Công việc", "Sản phẩm đầu ra"], [
    ("Ngày 1", "Phân tích đề, thiết kế dữ liệu và giao diện", "Schema, wireframe"),
    ("Ngày 2", "Khởi tạo SQLite, seed và API sản phẩm", "products CRUD API"),
    ("Ngày 3", "Trang chủ, tìm kiếm và giỏ hàng", "Home + cart"),
    ("Ngày 4", "Tạo hóa đơn, chi tiết và trừ tồn kho", "Order workflow"),
    ("Ngày 5", "Trang quản trị sản phẩm và hóa đơn", "Admin dashboard"),
    ("Ngày 6", "Responsive, kiểm thử và sửa lỗi", "Build đạt"),
    ("Ngày 7", "Ảnh kết quả, hướng dẫn và báo cáo", "Bộ hồ sơ nộp"),
], [0.9, 3.5, 2.1])
doc.add_heading("Rủi ro và cách xử lý", level=2)
bullet("Sai phiên bản môi trường: sử dụng `.nvmrc`, `.node-version`, `engines` và `packageManager`.")
bullet("Lỗi dữ liệu khi đặt hàng: kiểm tra đầu vào và dùng `BEGIN IMMEDIATE/COMMIT/ROLLBACK`.")
bullet("Chuyển máy bị lỗi ảnh: lưu ảnh trong thư mục `public/images`, không dùng URL ngoài.")

# 9
start_page("5. Công nghệ sử dụng", "Lựa chọn kỹ thuật phù hợp với thời gian và yêu cầu")
add_table(["Công nghệ", "Vai trò", "Lý do lựa chọn"], [
    ("Next.js 15", "Framework full-stack", "Trang và API trong cùng project"),
    ("React 19", "Giao diện tương tác", "Quản lý trạng thái giỏ, form và modal"),
    ("TypeScript", "Kiểm tra kiểu", "Giảm lỗi dữ liệu giữa UI, API và SQLite"),
    ("SQLite", "Cơ sở dữ liệu", "Một file, không cần cài DB server"),
    ("Lucide", "Biểu tượng", "Nhất quán, dễ nhận biết và nhẹ"),
    ("pnpm", "Quản lý thư viện", "Lockfile ổn định và tiết kiệm dung lượng"),
], [1.2, 1.65, 3.65])
doc.add_heading("Lưu ý về SQLite tích hợp", level=2)
para("Project sử dụng mô-đun `node:sqlite` có trong Node.js 22. Vì vậy máy chạy cần đúng Node 22 LTS như file cấu hình. Cách này loại bỏ dependency native bên ngoài và giúp quá trình cài đặt ít lỗi hơn giữa Windows, macOS và Linux.")

# 10
start_page("6. Kiến trúc hệ thống", "Mô hình ba lớp đơn giản")
add_picture(ASSETS / "architecture.png", width=5.0, caption="Hình 1. Kiến trúc tổng thể của TV House")
para("Lớp giao diện gồm các React component cho trang chủ, giỏ hàng và quản trị. Lớp API tiếp nhận JSON, kiểm tra dữ liệu và gọi hàm nghiệp vụ. Lớp dữ liệu trong `store-db.ts` đóng gói toàn bộ lệnh SQLite, giúp giao diện không truy cập database trực tiếp.")
bullet("Client chỉ gửi mã sản phẩm và số lượng; giá được đọc lại từ database.")
bullet("API trả JSON thống nhất, kèm mã HTTP phù hợp.")
bullet("SQLite bật khóa ngoại và chế độ WAL để nâng tính ổn định khi ghi dữ liệu.")

# 11
start_page("7. Tác nhân và ca sử dụng", "Các luồng nghiệp vụ chính")
add_table(["Tác nhân", "Ca sử dụng", "Kết quả"], [
    ("Khách hàng", "Tìm/lọc TV", "Danh sách phù hợp"), ("Khách hàng", "Thêm vào giỏ", "Giỏ lưu trên trình duyệt"),
    ("Khách hàng", "Đặt hàng", "Hóa đơn mới và tồn kho giảm"), ("Nhân viên", "Thêm sản phẩm", "Sản phẩm xuất hiện trên home"),
    ("Nhân viên", "Sửa sản phẩm", "Giá, tồn kho, mô tả được cập nhật"), ("Nhân viên", "Xóa sản phẩm", "Danh mục loại bỏ sản phẩm"),
    ("Nhân viên", "Xử lý hóa đơn", "Trạng thái được cập nhật"),
], [1.25, 2.35, 2.9])
doc.add_heading("Luồng đặt hàng tiêu chuẩn", level=2)
numbered("Khách chọn một hoặc nhiều TV từ trang chủ.")
numbered("Khách điều chỉnh số lượng và nhập thông tin giao hàng.")
numbered("Máy chủ kiểm tra sản phẩm, giá hiện tại và tồn kho.")
numbered("Transaction tạo orders, order_items và trừ stock.")
numbered("Trang thành công hiển thị mã hóa đơn cho khách.")

# 12
start_page("8. Thiết kế cơ sở dữ liệu", "Quan hệ giữa ba bảng")
add_table(["Bảng cha", "Quan hệ", "Bảng con", "Khóa ngoại"], [
    ("orders", "1 — N", "order_items", "order_items.order_id"),
    ("products", "1 — N", "order_items", "order_items.product_id"),
], [1.3, 1.1, 1.5, 2.6])
doc.add_heading("Nguyên tắc thiết kế", level=2)
bullet("Mỗi bảng dùng khóa chính số nguyên tự tăng để truy vấn đơn giản.")
bullet("`sku` và `code` là duy nhất, tránh trùng sản phẩm hoặc hóa đơn.")
bullet("`order_items` lưu `product_name` và `unit_price` để bảo toàn lịch sử.")
bullet("Xóa hóa đơn sẽ xóa chi tiết; xóa sản phẩm không xóa lịch sử hóa đơn.")
bullet("Các CHECK constraint ngăn giá âm, tồn kho âm và số lượng bằng 0.")
doc.add_heading("Chỉ mục", level=2)
para("Ba index được tạo theo truy vấn thực tế: lọc thương hiệu sản phẩm, sắp xếp hóa đơn theo ngày và lấy chi tiết theo `order_id`. Sau khi khởi tạo, `PRAGMA optimize` giúp SQLite cập nhật thông tin cho bộ lập kế hoạch truy vấn.")

# 13
start_page("9. Bảng products", "Lưu danh mục sản phẩm và tồn kho")
add_table(["Cột", "Kiểu", "Ràng buộc / ý nghĩa"], [
    ("id", "INTEGER", "PK, tự tăng"), ("sku", "TEXT", "Mã duy nhất"), ("name", "TEXT", "Tên TV"),
    ("brand", "TEXT", "Samsung/LG/Sony/TCL"), ("category", "TEXT", "LED/QLED/OLED..."),
    ("screen_size", "INTEGER", "Kích thước inch > 0"), ("resolution", "TEXT", "Độ phân giải"),
    ("price", "INTEGER", "Giá niêm yết ≥ 0"), ("sale_price", "INTEGER", "Có thể NULL"),
    ("stock", "INTEGER", "Tồn kho ≥ 0"), ("image", "TEXT", "Đường dẫn ảnh cục bộ"),
    ("description", "TEXT", "Mô tả sản phẩm"), ("featured", "INTEGER", "0 hoặc 1"),
], [1.35, 1.15, 4.0])
para("Giá tiền lưu bằng số nguyên đồng Việt Nam để tránh sai số số thực. Trường `updated_at` thay đổi khi sửa sản phẩm hoặc trừ tồn kho.")

# 14
start_page("10. Bảng orders và order_items", "Lưu hóa đơn cùng các dòng chi tiết")
doc.add_heading("10.1 Bảng orders", level=2)
add_table(["Cột", "Ý nghĩa"], [("code", "Mã hóa đơn dạng HDYYYYMMDD-XXXXX"), ("customer_name, phone, address", "Thông tin giao hàng"), ("payment_method", "COD hoặc BANK"), ("status", "Mới → Đang xử lý → Đang giao → Hoàn thành/Đã hủy"), ("total", "Tổng các subtotal")], [2.0, 4.5])
doc.add_heading("10.2 Bảng order_items", level=2)
add_table(["Cột", "Ý nghĩa"], [("order_id", "Khóa ngoại tới orders"), ("product_id", "Khóa ngoại có thể NULL sau khi xóa sản phẩm"), ("product_name", "Tên chụp tại lúc mua"), ("unit_price", "Đơn giá chụp tại lúc mua"), ("quantity", "Số lượng > 0"), ("subtotal", "unit_price × quantity")], [2.0, 4.5])

# 15
start_page("11. Thiết kế trang chủ", "Mặt tiền bán hàng của TV House")
bullet("Hero toàn màn hình giới thiệu TV và lời kêu gọi xem sản phẩm.")
bullet("Ba quyền lợi mua hàng: giao miễn phí, bảo hành và tư vấn.")
bullet("Thanh tìm kiếm theo tên, hãng hoặc công nghệ màn hình.")
bullet("Bộ lọc thương hiệu và lưới card sản phẩm responsive.")
bullet("Card hiển thị giảm giá, nhãn bán chạy, thông số, tồn kho và nút giỏ hàng.")
doc.add_heading("Quyết định trải nghiệm", level=2)
para("Giá khuyến mãi dùng màu đỏ để dễ nhận biết; thông tin còn hàng dùng xanh lá, sắp hết hàng dùng cam. Nút thêm giỏ là biểu tượng quen thuộc có tooltip. Kích thước card ổn định để nội dung không làm nhảy bố cục.")
doc.add_heading("Khả năng truy cập", level=2)
para("Ảnh có `alt`, nút icon có nhãn hoặc tooltip, form có `label`, điều hướng có `aria-label` và trạng thái active được thể hiện cả bằng màu lẫn đường gạch dưới.")

# 16
start_page("12. Thiết kế giỏ hàng", "Từ sản phẩm đến hóa đơn")
add_table(["Khu vực", "Chức năng"], [
    ("Danh sách sản phẩm", "Ảnh, tên, mã, giá, số lượng và nút xóa"),
    ("Bộ chỉnh số lượng", "Giới hạn từ 1 đến tồn kho hiện tại"),
    ("Thông tin giao hàng", "Họ tên, điện thoại và địa chỉ"),
    ("Thanh toán", "COD hoặc chuyển khoản"),
    ("Tóm tắt", "Tạm tính, phí giao hàng và tổng cộng"),
], [2.0, 4.5])
doc.add_heading("Quản lý trạng thái giỏ", level=2)
para("Giỏ hàng dùng `localStorage` vì đây là trạng thái tạm theo thiết bị, không phải nguồn dữ liệu chính thức. Sự kiện tùy chỉnh `tv-house-cart-change` đồng bộ số lượng trên header ngay khi thêm, xóa hoặc đặt hàng.")
doc.add_heading("Bảo vệ số tiền", level=2)
para("Tổng trên giao diện giúp khách tham khảo. Khi đặt hàng, phía máy chủ không tin giá gửi từ trình duyệt mà đọc lại sản phẩm trong SQLite, lấy `sale_price` hoặc `price`, sau đó tính tổng mới.")

# 17
start_page("13. Quản trị CRUD sản phẩm", "Thêm, sửa và xóa dữ liệu")
doc.add_heading("13.1 Danh sách và thống kê", level=2)
para("Dashboard hiển thị số sản phẩm, số hóa đơn, doanh thu chưa hủy và số mặt hàng có tồn kho dưới 8. Bảng sản phẩm cho phép tìm theo tên, SKU hoặc thương hiệu.")
doc.add_heading("13.2 Thêm và sửa", level=2)
para("Biểu mẫu dùng chung cho hai chế độ. Khi thêm, API sử dụng POST `/api/products`; khi sửa, API sử dụng PUT `/api/products/{id}`. Dữ liệu được chuẩn hóa SKU thành chữ hoa và kiểm tra các trường bắt buộc.")
doc.add_heading("13.3 Xóa", level=2)
para("Trước khi gọi DELETE, giao diện yêu cầu xác nhận. Khóa ngoại `ON DELETE SET NULL` cho phép xóa sản phẩm nhưng vẫn giữ tên, giá và số lượng trong hóa đơn cũ.")
add_table(["Thao tác", "HTTP", "Kết quả kiểm thử"], [("Thêm", "POST", "Tạo ID 9"), ("Sửa", "PUT", "Đổi tên, giá và tồn kho"), ("Xóa", "DELETE", "Trả {ok: true}, dữ liệu thử được dọn")], [1.5, 1.0, 4.0])

# 18
start_page("14. Quy trình tạo hóa đơn", "Transaction đảm bảo dữ liệu toàn vẹn")
numbered("Gộp các dòng cùng `productId` trong giỏ hàng.")
numbered("Đọc lại từng sản phẩm và kiểm tra tồn kho.")
numbered("Tính đơn giá từ `salePrice ?? price` và tính tổng hóa đơn.")
numbered("Mở `BEGIN IMMEDIATE` để giữ quyền ghi trong transaction.")
numbered("Insert bảng orders, sau đó insert từng order_items.")
numbered("Trừ stock với điều kiện `stock >= quantity`.")
numbered("COMMIT khi mọi bước đạt; ROLLBACK nếu có bất kỳ lỗi nào.")
doc.add_heading("Lợi ích", level=2)
para("Nếu hai khách đặt cùng sản phẩm gần như đồng thời, câu UPDATE có điều kiện ngăn tồn kho xuống âm. Nếu một dòng chi tiết thất bại, hóa đơn và mọi thay đổi tồn kho trong transaction cũng bị hủy, tránh dữ liệu nửa chừng.")
doc.add_heading("Dữ liệu mẫu sau kiểm thử", level=2)
add_table(["Mã hóa đơn", "Khách", "Số dòng", "Tổng"], [("HD20260825-KDHFU", "Nguyễn Văn An", "2", "42.470.000đ")], [2.0, 1.8, 1.0, 1.7])

# 19
start_page("15. Thiết kế API", "Các endpoint dùng trong chương trình")
add_table(["Method", "Đường dẫn", "Chức năng"], [
    ("GET", "/api/products", "Danh sách sản phẩm"), ("POST", "/api/products", "Thêm sản phẩm"),
    ("GET", "/api/products/{id}", "Chi tiết sản phẩm"), ("PUT", "/api/products/{id}", "Cập nhật sản phẩm"),
    ("DELETE", "/api/products/{id}", "Xóa sản phẩm"), ("GET", "/api/orders", "Danh sách hóa đơn"),
    ("POST", "/api/orders", "Đặt hàng"), ("GET", "/api/orders/{id}", "Hóa đơn + chi tiết"),
    ("PATCH", "/api/orders/{id}", "Đổi trạng thái"),
], [0.8, 2.3, 3.4])
doc.add_heading("Quy ước phản hồi", level=2)
bullet("Thành công trả object `product`, `products`, `order` hoặc `orders`.")
bullet("Tạo mới trả HTTP 201; không tìm thấy trả 404; dữ liệu sai trả 400.")
bullet("Lỗi được chuyển thành thông báo tiếng Việt để giao diện hiển thị trực tiếp.")

# 20
start_page("16. Đoạn mã chính 1", "Khởi tạo cấu trúc bảng sản phẩm")
add_picture(ASSETS / "code-schema.png", width=5.8, caption="Hình 2. Mã tạo bảng products trong store-db.ts")
para("Đoạn mã dùng `CREATE TABLE IF NOT EXISTS` nên có thể chạy lại nhiều lần mà không mất dữ liệu. Chế độ STRICT giúp SQLite kiểm tra kiểu dữ liệu chặt hơn. Các CHECK constraint bảo đảm kích thước, giá và tồn kho hợp lệ; UNIQUE trên SKU ngăn trùng mã sản phẩm.")

# 21
start_page("17. Đoạn mã chính 2", "Transaction khi tạo hóa đơn")
add_picture(ASSETS / "code-transaction.png", width=5.8, caption="Hình 3. Mã transaction trong hàm createOrder")
para("Transaction bắt đầu bằng `BEGIN IMMEDIATE`. Hóa đơn được insert trước để lấy `lastInsertRowid`, sau đó từng chi tiết được lưu và tồn kho được trừ có điều kiện. COMMIT chỉ xảy ra khi tất cả bước hoàn tất; khối catch thực hiện ROLLBACK rồi chuyển lỗi cho API.")

# 22
start_page("18. Đoạn mã chính 3", "API tiếp nhận và kiểm tra đơn hàng")
add_picture(ASSETS / "code-api-order.png", width=5.8, caption="Hình 4. API POST /api/orders")
para("API chuẩn hóa dữ liệu JSON, chỉ nhận hai phương thức COD/BANK, chuyển productId và quantity thành số nguyên, sau đó kiểm tra tên, số điện thoại, địa chỉ và giỏ hàng. Chỉ dữ liệu hợp lệ mới được chuyển sang hàm nghiệp vụ `createOrder`.")

# 23
start_page("19. Kế hoạch kiểm thử", "Kiểm thử tĩnh, tích hợp và giao diện")
add_table(["Nhóm", "Ca kiểm thử", "Kỳ vọng"], [
    ("Build", "typecheck, lint, next build", "Không có lỗi"),
    ("Trang", "GET /, /cart, /admin", "HTTP 200"),
    ("Database", "Khởi động khi chưa có file", "Tạo 3 bảng + seed"),
    ("Order", "Đặt 2 sản phẩm", "Tạo orders + 2 order_items"),
    ("Stock", "Sau khi đặt", "Giảm đúng 1 và 2"),
    ("CRUD", "Thêm → sửa → xóa sản phẩm tạm", "Cả 3 API thành công"),
    ("Responsive", "Viewport 390 × 844", "Không tràn ngang"),
], [1.1, 3.1, 2.3])
doc.add_heading("Kết quả tổng quát", level=2)
bullet("TypeScript và ESLint: đạt.")
bullet("Next.js production build: đạt, 21 trang/API được tạo.")
bullet("Trang chủ và quản trị: HTTP 200.")
bullet("CRUD và tạo hóa đơn: đạt theo dữ liệu thực tế.")
bullet("Mobile: `bodyScrollWidth` không vượt `innerWidth`.")

# 24
start_page("20. Kết quả chạy – Trang chủ", "Giao diện bán hàng trên màn hình desktop")
add_picture(ASSETS / "result-home.png", width=6.3, caption="Hình 5. Trang chủ TV House với hero, quyền lợi và đầu danh mục sản phẩm")
para("Trang chủ hiển thị đúng tên thương hiệu, ảnh TV cục bộ, các nút điều hướng và 8 sản phẩm mẫu. Giá, mức giảm, tồn kho và bộ lọc thương hiệu được lấy từ SQLite thông qua server component.")

# 25
start_page("21. Kết quả chạy – Giỏ hàng", "Sản phẩm, số lượng và biểu mẫu giao hàng")
add_picture(ASSETS / "result-cart.png", width=6.3, caption="Hình 6. Trang giỏ hàng với một TV và tổng thanh toán")
para("Sau khi bấm biểu tượng giỏ trên card, badge header đổi thành 1. Trang giỏ hiển thị đúng sản phẩm, giá khuyến mãi, bộ tăng giảm số lượng, hai phương thức thanh toán và tổng cộng.")

# 26
start_page("22. Kết quả chạy – Quản trị", "Danh sách hóa đơn và số liệu tổng quan")
add_picture(DOCS / "screenshots/03-quan-tri-hoa-don.png", width=6.3, caption="Hình 7. Dashboard quản trị sau khi tạo hóa đơn mẫu")
para("Dashboard ghi nhận 8 sản phẩm, 1 hóa đơn, doanh thu 42.470.000đ và 3 sản phẩm sắp hết hàng. Bảng hóa đơn hiển thị khách hàng, thời gian, tổng tiền, combobox trạng thái và nút xem chi tiết.")

# 27
start_page("23. Hướng dẫn triển khai", "Chạy project trên máy tính khác")
doc.add_heading("23.1 Các bước", level=2)
numbered("Cài Node.js 22 LTS và mở Terminal mới.")
numbered("Chạy `corepack enable` và kích hoạt pnpm 9.12.0.")
numbered("Mở thư mục có package.json, chạy `pnpm install --frozen-lockfile`.")
numbered("Chạy `pnpm run dev`, mở http://localhost:3005.")
numbered("Để kiểm tra production, chạy typecheck, lint, test, build và start.")
doc.add_heading("23.2 Database", level=2)
para("SQLite nằm tại `data/tv-store.db`. Nếu file chưa tồn tại, project tự tạo bảng và 8 sản phẩm mẫu. Khi sao lưu phải dừng ứng dụng và chép file `.db`; không chép file `-wal` hoặc `-shm` đang hoạt động.")
doc.add_heading("23.3 Lưu ý triển khai", level=2)
para("Mô hình file SQLite phù hợp một máy, phòng máy hoặc VPS có ổ đĩa bền vững. Không dùng trực tiếp trên hosting serverless có hệ thống file tạm. Hướng dẫn đầy đủ nằm trong `docs/HUONG_DAN_TRIEN_KHAI.md`.")

# 28
start_page("24. Mã nguồn và cấu trúc", "Thông tin để nộp link code")
add_table(["Thành phần", "Vị trí"], [
    ("Mã nguồn chính", "src/"), ("SQLite", "data/tv-store.db và src/lib/store-db.ts"),
    ("Schema SQL", "database/schema.sql"), ("Ảnh kết quả", "docs/screenshots/"),
    ("Hướng dẫn", "README.md và docs/HUONG_DAN_TRIEN_KHAI.md"), ("Gói nộp", "TV-House-Source.zip"),
], [2.0, 4.5])
doc.add_heading("Liên kết mã nguồn", level=2)
para("GitHub / Google Drive: [DÁN LIÊN KẾT CHIA SẺ TẠI ĐÂY]")
doc.add_heading("Các file chính", level=2)
bullet("`src/lib/store-db.ts`: schema, seed, truy vấn và transaction.")
bullet("`src/components/store/product-catalog.tsx`: tìm kiếm, lọc và thêm giỏ.")
bullet("`src/components/store/cart-page.tsx`: giỏ hàng và checkout.")
bullet("`src/components/store/admin-dashboard.tsx`: CRUD và hóa đơn.")
bullet("`src/app/api/`: các endpoint JSON.")
para("Trước khi nộp, sinh viên tải gói ZIP lên nơi lưu trữ do giảng viên yêu cầu, bật quyền xem và thay placeholder phía trên bằng liên kết thực tế.")

# 29
start_page("25. Kết luận và hướng phát triển", "Đánh giá kết quả sau một tuần")
para("Đề tài đã đáp ứng các yêu cầu cốt lõi: cơ sở dữ liệu đúng ba bảng, trang home, trang giỏ hàng, chức năng thêm/sửa/xóa sản phẩm, tạo hóa đơn và chi tiết hóa đơn. Project có thể cài lại trên máy khác nhờ phiên bản môi trường cố định, lockfile và tài liệu triển khai.")
doc.add_heading("Hướng phát triển", level=2)
bullet("Bổ sung đăng nhập và phân quyền quản trị.")
bullet("Tải nhiều ảnh sản phẩm và tạo trang chi tiết TV.")
bullet("Tích hợp cổng thanh toán và thông báo email.")
bullet("Thêm quản lý khách hàng, khuyến mãi và báo cáo doanh thu theo thời gian.")
bullet("Dùng migration có phiên bản khi schema phát triển lớn hơn.")
doc.add_heading("Tài liệu tham khảo", level=2)
numbered("Next.js Documentation — App Router và Route Handlers.")
numbered("Node.js Documentation — mô-đun node:sqlite.")
numbered("SQLite Documentation — foreign keys, transactions và indexes.")
numbered("React Documentation — state, event và component composition.")
numbered("TypeScript Handbook — static typing cho ứng dụng web.")
para("— HẾT —", align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

doc.core_properties.title = "Báo cáo giữa kỳ - Website bán tivi TV House"
doc.core_properties.subject = "Next.js, TypeScript và SQLite"
doc.core_properties.author = "Sinh viên - điền trước khi nộp"
doc.core_properties.keywords = "TV House, Next.js, SQLite, giữa kỳ"
doc.save(OUTPUT)
print(OUTPUT)
